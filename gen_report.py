"""生成项目进度报告 Word 文档（约 75% 完成度，无 emoji）"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 标题 ──
title = doc.add_heading('NCF 个性化推荐系统 -- 项目进度报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('项目路径：ncf_reproduction')
doc.add_paragraph('报告日期：2026年6月17日')
doc.add_paragraph('整体状态：核心框架已搭建完成，部分模块仍在开发中，整体进度约 75%')

doc.add_paragraph()

# ── 1. 项目整体架构 ──
doc.add_heading('一、项目整体架构', level=1)
doc.add_paragraph(
    '本项目是一个基于 PyTorch 的个性化推荐系统，采用「多路召回 + 深度学习排序」'
    '的两阶段架构，支持 MovieLens-1M 和 Last.fm 两个公开数据集。'
    '系统由下至上分为四层：数据层、召回层、排序层、服务层。'
)

# 架构表格
arch_table = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ('服务层', 'Gradio 可视化界面 + FastAPI RESTful 服务'),
    ('排序层', 'GMF / MLP / NeuMF / Wide&Deep 深度学习精排'),
    ('召回层', 'UserCF / ItemCF / SVD / NMF 多路召回融合'),
    ('数据层', 'BaseDataset 抽象 -> MovieLens-1M / Last.fm'),
    ('评估', 'HR@K / NDCG@K / MAP / Diversity / Coverage / Novelty'),
]
for i, (k, v) in enumerate(cells):
    arch_table.rows[i].cells[0].text = k
    arch_table.rows[i].cells[1].text = v

doc.add_paragraph()

# ── 2. 各模块进度 ──
doc.add_heading('二、各模块进度明细', level=1)

# ── 模块进度汇总表 ──
doc.add_heading('2.1 进度总览', level=2)
summary_table = doc.add_table(rows=10, cols=3, style='Light Shading Accent 1')
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['模块', '状态', '完成度']):
    summary_table.rows[0].cells[i].text = h
    for p in summary_table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
summary_data = [
    ('数据层 (data_loader.py)', '已完成', '100%'),
    ('召回层 (recall/)', '基本完成，NMF 待调优', '85%'),
    ('排序层 (model.py + rank/)', 'GMF/MLP/Wide&Deep 已完成，NeuMF 预训练待调试', '75%'),
    ('推荐管道 (pipeline.py)', '已完成', '100%'),
    ('冷启动 (cold_start.py)', '基础框架完成，内容特征融合待实现', '60%'),
    ('评估系统 (evaluate.py)', 'HR/NDCG 已完成，扩展指标部分完成', '70%'),
    ('服务层 (api/)', 'FastAPI 已完成，Gradio 界面待完善', '65%'),
    ('实验系统 (experiments/)', '对比实验已完成，冷启动实验待跑', '60%'),
    ('小论文 (paper/)', '初稿已撰写，实验数据待补充完整', '50%'),
]
for r, row_data in enumerate(summary_data):
    for c, val in enumerate(row_data):
        summary_table.rows[r + 1].cells[c].text = val

doc.add_paragraph()

# ── 2.2 数据层 ──
doc.add_heading('2.2 数据层 -- data_loader.py [已完成]', level=2)
items = [
    'BaseDataset 抽象基类：统一数据集接口（download / load / split / neg_sample）',
    'MovieLensDataset：自动下载并加载 MovieLens-1M（6040 用户 x 3706 电影 x 100 万评分）',
    'LastfmDataset：支持 Last.fm HetRec 2011 数据集（1892 用户 x 17632 艺术家）',
    'Leave-One-Out 数据划分：最新交互->测试集，倒数第二->验证集，其余->训练集',
    '负采样（4:1 正负比）+ 困难负采样（hard_negative_sampling）',
    '冷启动用户分组工具（cold_start_split）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ── 2.3 召回层 ──
doc.add_heading('2.3 召回层 -- recall/ [基本完成]', level=2)
recall_table = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
recall_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['模型', '文件', '方法', '状态']
for i, h in enumerate(headers):
    recall_table.rows[0].cells[i].text = h
recall_data = [
    ('UserCF', 'cf.py', '基于用户的协同过滤，余弦相似度，稀疏矩阵加速', '已完成'),
    ('ItemCF', 'cf.py', '基于物品的协同过滤', '已完成'),
    ('SVD', 'svd.py', '含偏置矩阵分解，PyTorch 实现，BCE Loss', '已完成'),
    ('NMF', 'nmf.py', '非负矩阵分解，exp 激活保证非负性', '待调优'),
]
for r, row_data in enumerate(recall_data):
    for c, val in enumerate(row_data):
        recall_table.rows[r + 1].cells[c].text = val
doc.add_paragraph('统一 RecallBase 抽象接口已完成，多路召回融合（取并集去重）逻辑已实现。NMF 召回效果较差（HR@10 仅 0.0137），需要进一步调优训练策略。')

# ── 2.4 排序层 ──
doc.add_heading('2.4 排序层 -- model.py + rank/ [大部分完成]', level=2)
rank_table = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
rank_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    rank_table.rows[0].cells[i].text = h
rank_data = [
    ('GMF', 'model.py', '嵌入逐元素乘积 + 线性输出 + Sigmoid', '已完成'),
    ('MLP', 'model.py', '嵌入拼接 -> 塔式 MLP [32,16,8] -> Sigmoid', '已完成'),
    ('NeuMF', 'model.py', 'GMF+MLP 双分支融合，支持预训练 + Alpha 融合', '预训练待调试'),
    ('Wide&Deep', 'rank/wide_deep.py', 'Wide(交叉特征) + Deep(MLP) + Dropout', '已完成'),
]
for r, row_data in enumerate(rank_data):
    for c, val in enumerate(row_data):
        rank_table.rows[r + 1].cells[c].text = val
doc.add_paragraph('统一 BaseRanker 抽象接口已完成。GMF、MLP、Wide&Deep 三个模型训练正常。NeuMF 随机初始化模式可用，但预训练模式（先训 GMF 和 MLP 再加载权重融合）的训练流程较长，收敛效果待验证。')

# ── 2.5 管道 ──
doc.add_heading('2.5 两阶段推荐管道 -- pipeline.py [已完成]', level=2)
doc.add_paragraph('RecommendationPipeline 类已实现：多路召回 -> 融合去重 -> 排序 -> Top-N。支持来源追踪（标注每个推荐结果来自哪个召回方法）。')

# ── 2.6 冷启动 ──
doc.add_heading('2.6 冷启动处理 -- cold_start.py [部分完成]', level=2)
doc.add_paragraph('ColdStartHandler 类已实现基础框架：四级用户分群（极度冷启动 <=1 / 冷启动 2-3 / 温暖 4-5 / 热门 6+），基于热度的差异化推荐策略，可调节多样性注入权重。')
doc.add_paragraph('待完成：引入电影类型（genre）等内容特征辅助冷启动推荐，目前仅依赖全局热度统计，对偏好小众品类的用户效果有限。')

# ── 2.7 评估系统 ──
doc.add_heading('2.7 评估系统 -- evaluate.py [部分完成]', level=2)
items2 = [
    '基础指标已完成：hit_ratio_at_k / ndcg_at_k（负采样评估，99 负样本，rank among 100）',
    '扩展指标部分完成：map_at_k 已实现；diversity_at_k / coverage_at_k / novelty_at_k 已实现但未在实验中系统使用',
    '工具函数：compute_item_popularity / compute_item_similarity 已完成',
    '待完成：full_evaluate 与训练流程的集成，指标可视化导出',
]
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

# ── 2.8 服务部署 ──
doc.add_heading('2.8 服务部署 -- api/ [部分完成]', level=2)
items3 = [
    'FastAPI 服务 (server.py) 已完成：/recommend/{user_id}、/recommend/cold-start、/health、/models 四个端点',
    'Gradio 界面 (app.py) 基础框架已完成：三页签布局（Top-N 推荐 / 模型对比 / 冷启动测试）',
    'Gradio 界面待完善：模型对比图表目前需手动传入 metrics 数据，缺少自动加载实验结果的功能',
    '一键启动 (launch.py) 已完成：自动训练召回模型并启动 Gradio，但排序模型未集成（ranker=None）',
]
for item in items3:
    doc.add_paragraph(item, style='List Bullet')

# ── 2.9 实验系统 ──
doc.add_heading('2.9 实验系统 -- experiments/run_experiments.py [部分完成]', level=2)
doc.add_paragraph('实验 1（模型对比）已完成：传统方法（ItemCF / SVD / NMF）vs 深度方法（GMF / MLP / NeuMF / Wide&Deep）共 6 个模型的 HR@10 和 NDCG@10 对比。')
doc.add_paragraph('实验 2（冷启动分析）代码框架已完成，但尚未系统运行和整理结果数据。')
doc.add_paragraph('支持 --dataset movielens / lastfm 切换数据集，但 Last.fm 数据集上的完整实验尚未进行。')

# ── 2.10 论文 ──
doc.add_heading('2.10 小论文 -- paper/小论文.md [初稿阶段]', level=2)
doc.add_paragraph('中文论文初稿约 290 行，结构完整（摘要、引言、研究现状、理论技术、方法、实验、结束语、参考文献 15 篇），部分实验数据已填充，但冷启动实验和 Last.fm 实验数据尚待补充。英文摘要中还有占位标记待替换。')

doc.add_paragraph()

# ── 3. 已有实验结果 ──
doc.add_heading('三、已有实验结果（MovieLens-1M，初步）', level=1)
doc.add_paragraph('以下为模型对比实验的初步结果，部分模型数据可能随后续调优更新：')

exp_table = doc.add_table(rows=8, cols=3, style='Light Shading Accent 1')
exp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
exp_headers = ['模型', 'HR@10', 'NDCG@10']
for i, h in enumerate(exp_headers):
    exp_table.rows[0].cells[i].text = h
    for paragraph in exp_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
exp_data = [
    ('ItemCF', '0.0692', '0.0351'),
    ('SVD', '0.0268', '0.0124'),
    ('NMF', '0.0137', '0.0061'),
    ('GMF', '0.6210', '0.3495'),
    ('MLP', '0.6116', '0.3420'),
    ('NeuMF', '0.6386', '0.3689'),
    ('Wide&Deep (最优)', '0.6517', '0.3766'),
]
for r, row_data in enumerate(exp_data):
    for c, val in enumerate(row_data):
        exp_table.rows[r + 1].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('初步结论：')
run.bold = True
p.add_run(
    '深度学习方法（GMF/MLP/NeuMF/Wide&Deep）的 HR@10 均在 0.61 以上，'
    '显著优于传统协同过滤方法（均低于 0.07）。Wide&Deep 模型表现最优'
    '（HR@10=0.6517, NDCG@10=0.3766）。NeuMF 通过融合 GMF 和 MLP 分支'
    '（HR@10=0.6386）优于单一分支模型，验证了多分支融合架构的有效性。'
    '注：以上结果基于当前训练配置，后续调优后可能变化。'
)

doc.add_paragraph()

# ── 4. 文件清单 ──
doc.add_heading('四、文件清单', level=1)
doc.add_paragraph('共 20 个 Python 源文件：')
file_list = [
    'model.py -- GMF / MLP / NeuMF 模型定义',
    'train.py -- 训练循环 + 早停机制',
    'evaluate.py -- HR / NDCG / MAP / Diversity / Coverage / Novelty',
    'main.py -- 命令行入口（NCF 训练）',
    'data_loader.py -- 数据加载 / 划分 / 负采样 / BaseDataset',
    'pipeline.py -- 两阶段推荐管道',
    'cold_start.py -- 冷启动处理',
    'launch.py -- 一键启动 Gradio',
    'recall/base.py -- 召回抽象基类',
    'recall/cf.py -- UserCF + ItemCF',
    'recall/svd.py -- SVD 矩阵分解召回',
    'recall/nmf.py -- NMF 非负矩阵分解召回',
    'rank/base.py -- 排序抽象基类',
    'rank/wide_deep.py -- Wide&Deep 模型',
    'api/app.py -- Gradio 可视化界面',
    'api/server.py -- FastAPI 服务',
    'experiments/run_experiments.py -- 对比实验 + 冷启动实验',
]
for f in file_list:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()

# ── 5. 后续工作计划 ──
doc.add_heading('五、后续工作计划', level=1)

doc.add_heading('近期待完成（优先级高）', level=2)
todo_high = [
    '完成 NeuMF 预训练模式的调试，验证 Alpha 融合权重对效果的影响',
    '运行并整理冷启动分组实验，补充论文中冷启动场景的数据',
    '在 Last.fm 数据集上运行完整实验，获得跨数据集的对比结论',
    '补充论文中英文摘要的占位数据，完善实验分析章节',
]
for item in todo_high:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('中期计划（优先级中）', level=2)
todo_mid = [
    'NMF 召回模型调优，改进负采样策略或尝试 BPR 损失函数',
    '将排序模型（Wide&Deep）集成到 launch.py 和 Gradio 界面中',
    'Gradio 界面完善：自动加载实验结果、模型指标可视化',
    '引入电影 genre 类型信息辅助冷启动推荐',
]
for item in todo_mid:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('长期优化（优先级低）', level=2)
todo_low = [
    '引入注意力机制学习用户交互序列中的时序模式',
    '添加单元测试覆盖核心模块',
    '模型权重保存与加载机制，支持推理时直接加载预训练模型',
    '考虑引入更多排序模型（DeepFM、DCN 等）进行对比',
]
for item in todo_low:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# ── 总体进度 ──
doc.add_heading('总体进度', level=1)
p = doc.add_paragraph()
run = p.add_run('进度评估：约 75% 完成')
run.bold = True
run.font.size = Pt(14)
doc.add_paragraph(
    '项目核心框架已搭建完成，数据层、召回层（除 NMF 调优外）、排序层（除 NeuMF 预训练外）'
    '和推荐管道已可正常运行。FastAPI 服务和 Gradio 界面基础功能可用。'
    '主要差距在于：冷启动实验数据待补充、Last.fm 跨数据集验证待完成、'
    '论文实验章节待完善、部分模块的集成与调优工作仍在进行中。'
    '预计在完成上述近期待办事项后，进度可提升至 90% 以上。'
)

# ── 保存 ──
output_path = r'c:\Users\黄海亦\Desktop\ncf_reproduction\项目进度报告_75.docx'
doc.save(output_path)
print(f'报告已生成：{output_path}')
